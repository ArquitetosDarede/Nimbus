"""
DOCX Analysis Agent - Specialized agent for DOCX file analysis
"""

import logging
from typing import Dict, Any
from pathlib import Path

logger = logging.getLogger(__name__)

# Optional imports
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DocxAnalysisAgent:
    """Specialized agent for DOCX file analysis."""

    def analyze(self, file_path: Path) -> Dict[str, Any]:
        """Extract text and tables from DOCX files."""
        if not HAS_DOCX:
            raise ImportError("python-docx not available for DOCX extraction")

        doc = Document(file_path)
        text_content = []
        tables = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            headers = []

            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)

                if row_idx == 0:  # First row as headers
                    headers = row_data
                table_data.append(row_data)

            if table_data:
                tables.append({
                    "table_index": table_idx,
                    "data": table_data,
                    "headers": headers
                })

        return {
            "extracted_text": "\n\n".join(text_content),
            "extracted_data": {
                "tables": tables,
                "metadata": {
                    "paragraphs": len(text_content),
                    "tables": len(tables)
                }
            }
        }